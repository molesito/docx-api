from flask import Flask, request, jsonify, send_file
from docx import Document
import io
import os

app = Flask(__name__)

def replace_placeholders(doc: Document, campos: dict):
    """
    Busca en todo el documento los placeholders [key] y los reemplaza por los valores.
    """
    for p in doc.paragraphs:
        for key, value in campos.items():
            placeholder = f"[{key}]"
            if placeholder in p.text:
                inline = p.runs
                for run in inline:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in campos.items():
                    placeholder = f"[{key}]"
                    if placeholder in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

@app.route("/fill-doc", methods=["POST"])
def fill_doc():
    try:
        # Recibir archivo y campos
        file = request.files.get("file")
        campos = request.json.get("campos")

        if not file or not campos:
            return jsonify({"error": "Falta el archivo .docx o el campo 'campos'"}), 400

        # Abrir docx en memoria
        doc = Document(file)

        # Procesar cada diccionario en el array de campos
        for entry in campos:
            replace_placeholders(doc, entry)

        # Guardar resultado en memoria
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name="resultado.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
